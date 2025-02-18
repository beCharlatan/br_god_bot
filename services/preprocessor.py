import argparse
from pathlib import Path
from docx import Document
from .document_embedder import DocumentEmbedder
from typing import List, Dict

class Preprocessor:
    def __init__(self):
        self.embedder = DocumentEmbedder()
    
    def parse_arguments(self) -> argparse.Namespace:
        parser = argparse.ArgumentParser(description='Генерация тест-кейсов по бизнес-требованниям')
        parser.add_argument('input_file', type=str, help='Path to DOCX file with business requirements')
        return parser.parse_args()
    
    def validate_file(self, file_path: str) -> None:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        if path.suffix.lower() != '.docx':
            raise ValueError(f"Invalid file type. Expected .docx file, got: {path.suffix}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            return "\n\n".join(paragraphs)
        except Exception as e:
            raise ValueError(f"Error reading docx file: {str(e)}")
    
    def find_similar_documents(self, query_text: str, limit: int = 5) -> List[Dict]:
        return self.embedder.find_similar_documents(query_text, limit)
    
    def read_requirements_file(self, file_path: str) -> str:
        self.validate_file(file_path)
        content = self.extract_text_from_docx(file_path)
        
        self.embedder.process_document(file_path, content)
        
        return content
