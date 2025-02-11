from pathlib import Path
from docx import Document

class Preprocessor:
    def validate_file(self, file_path: str) -> None:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        if path.suffix.lower() != '.docx':
            raise ValueError(f"Invalid file type. Expected .docx file, got: {path.suffix}")
    
    def extract_text_from_docx(self, file_path: str) -> tuple[str, str]:
        """Extract text content and main title from a DOCX file.
        
        Returns:
            tuple: (main_title, content)
            - main_title: The first heading found in the document
            - content: Full document content
        """
        try:
            doc = Document(file_path)
            paragraphs = []
            main_title = None
            
            # Find the first heading (main title)
            for paragraph in doc.paragraphs:
                if not main_title and paragraph.style.name.startswith('Heading'):
                    main_title = paragraph.text.strip()
                    break
            
            # If no heading found, use first non-empty paragraph
            if not main_title:
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        main_title = paragraph.text.strip()
                        break
            
            # Get all paragraphs and table content
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            content = "\n\n".join(paragraphs)
            return main_title or "Untitled Document", content
            
        except Exception as e:
            raise ValueError(f"Error reading docx file: {str(e)}")

    def read_requirements_file(self, file_path: str) -> tuple[str, str]:
        """Read requirements file and extract title and content.
        
        Returns:
            tuple: (title, content)
            - title: The main title of the document
            - content: Full document content
        """
        self.validate_file(file_path)
        return self.extract_text_from_docx(file_path)
